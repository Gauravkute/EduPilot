from docx import Document

def extract(filepath: str) -> dict:
    doc = Document(filepath)

    #Paragraph
    paragraphs =[]
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        paragraphs.append({
            "style":para.style.name,
            "text":para.text,
            "bold":any(run.bold for run in para.runs),
            "italic" : any(run.italic for run in para.runs),
        })

    #Tables
    tables =[]
    for table in doc.tables:
        grid = []
        for row in table.rows:
            grid.append([cell.text.strip() for cell in row.cells])
        tables.append(grid)

    #Metadata
    props = doc.core_properties
    meta ={
        "author":props.author,
        "title":props.title,
        "created":str(props.created),
    }

    return {"paragraph":paragraphs,"tables":tables,"metadata":meta}
